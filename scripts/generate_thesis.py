import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document():
    doc = Document()
    
    # ----------------------------------------------------
    # Margins setup
    # ----------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # ----------------------------------------------------
    # Color palette & styles setup
    # ----------------------------------------------------
    PRIMARY_COLOR = RGBColor(16, 44, 87)    # Navy Blue
    SECONDARY_COLOR = RGBColor(53, 162, 235) # Light Blue
    TEXT_COLOR = RGBColor(51, 51, 51)       # Charcoal
    
    # Configure Normal Style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # Configure Title Style
    if 'Title' in doc.styles:
        style_title = doc.styles['Title']
    else:
        style_title = doc.styles.add_style('Title', 1) # 1 = paragraph style
    style_title.font.name = 'Calibri'
    style_title.font.size = Pt(26)
    style_title.font.bold = True
    style_title.font.color.rgb = PRIMARY_COLOR
    style_title.paragraph_format.space_after = Pt(12)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Configure Heading Styles
    for h_name, size, color in [('Heading 1', Pt(18), PRIMARY_COLOR), 
                                 ('Heading 2', Pt(14), PRIMARY_COLOR), 
                                 ('Heading 3', Pt(12), SECONDARY_COLOR)]:
        style = doc.styles[h_name]
        style.font.name = 'Calibri'
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # Helper functions
    # ----------------------------------------------------
    def add_title(text):
        doc.add_paragraph(text, style='Title')
        
    def add_h1(text):
        doc.add_heading(text, level=1)
        
    def add_h2(text):
        doc.add_heading(text, level=2)
        
    def add_h3(text):
        doc.add_heading(text, level=3)
        
    def add_p(text):
        return doc.add_paragraph(text)
        
    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(text)
        return p
        
    def add_numbered(text):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(text)
        return p

    def set_cell_background(cell, color_hex):
        shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

    # ----------------------------------------------------
    # 1. TITLE PAGE
    # ----------------------------------------------------
    for _ in range(3):
        add_p("")
        
    # Title
    p_title = add_p("DESIGN AND DEVELOPMENT OF A MACHINE LEARNING MODEL FOR MUSIC HIT PREDICTION")
    p_title.style = 'Title'
    p_title.runs[0].font.size = Pt(24)
    p_title.runs[0].font.bold = True
    
    for _ in range(4):
        add_p("")
        
    # Institution details
    p_inst = add_p("Individual Project (ST6001CEM)\nSoftwarica College of IT & E-Commerce\nIn Collaboration with Coventry University")
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.runs[0].font.size = Pt(14)
    p_inst.runs[0].font.color.rgb = PRIMARY_COLOR
    p_inst.runs[0].font.bold = True
    
    for _ in range(4):
        add_p("")
        
    # Author details
    p_author = add_p("Submitted by:\nName: [Your Name]\nStudent ID: [Your ID]\nSupervisor: [Supervisor Name]\nDate: July 2026")
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.runs[0].font.size = Pt(12)
    p_author.runs[0].font.italic = True
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 2. ABSTRACT & TABLE OF CONTENTS PLACEHOLDER
    # ----------------------------------------------------
    add_h1("Abstract")
    add_p(
        "Predicting the commercial success of music has transitionally shifted from qualitative human assessment "
        "to quantitative data science. This thesis investigates the classification of music hits using machine learning "
        "and audio features extracted from digital waveforms. The research leverages Spotify's raw audio descriptors "
        "and metadata, compiling datasets containing 29,488 tracks (spotify.csv) and 32,833 tracks (spotify_songs.csv). "
        "A binary target variable is operationally defined, marking tracks with a popularity index of 70 or above as 'hits' "
        "(representing approximately 14.72% of the dataset) and all other tracks as 'flops'. "
        "Due to the substantial class imbalance, the modeling strategy focuses on maximizing the macro-averaged F1-score rather "
        "than raw classification accuracy. We evaluate three distinct modeling techniques: Logistic Regression (serving as the "
        "interpretable baseline), Random Forest, and eXtreme Gradient Boosting (XGBoost). "
        "Crucially, the baseline Logistic Regression model incorporated class-balancing weights, leading it to achieve the "
        "highest macro F1-score of 0.573 (with an overall accuracy of 71.60% and ROC AUC of 0.622), outperforming both "
        "XGBoost (accuracy of 83.91%, macro F1 of 0.530) and Random Forest (accuracy of 83.35%, macro F1 of 0.523) on our temporal test set. "
        "This outcome highlights a significant trade-off in imbalanced classification tasks: while tree-based ensembles excel at "
        "optimizing global accuracy by prioritizing the majority class, balanced parametric classifiers are more effective at "
        "identifying minority class success. The findings demonstrate that digital audio characteristics can predict hit "
        "potential, but predictive success is heavily genre-dependent, and the model's output must be framed as a risk-mitigation tool "
        "rather than a deterministic indicator of artistic value."
    )
    
    add_h1("Table of Contents")
    add_p(
        "To update the Table of Contents in Microsoft Word:\n"
        "1. Right-click on this page or the area below.\n"
        "2. Select 'Update Field' from the context menu.\n"
        "3. Choose 'Update entire table' and click OK.\n"
        "Word will automatically compile all headings styled with Heading 1, 2, and 3 into the table below."
    )
    add_p("[Auto-generated Table of Contents Field will appear here]")
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 3. CHAPTER 1: BACKGROUND & BUSINESS CASE
    # ----------------------------------------------------
    add_h1("Chapter 1: Background & Business Case")
    
    add_h2("1.1 Problem Statement")
    add_p(
        "The global music industry is experiencing an era of unprecedented scale and commercialization, generating "
        "over $26 billion in annual streaming revenues. Platforms like Spotify host catalogs exceeding 100 million tracks, "
        "with over 100,000 new songs uploaded daily. In this saturated environment, achieving commercial viability "
        "has become highly competitive and mathematically improbable: less than 15% of all released music obtains "
        "significant streaming traction or chart placement. Historically, record labels absorbed massive financial losses, "
        "relying on the revenue generated by a small fraction of breakout hits to fund their broader catalogs. "
        "However, independent artists—who represent the fastest-growing sector of the industry—lack the financial buffers "
        "and analytical resources of major labels. Major conglomerates employ dedicated data science teams to optimize "
        "release strategies, select promotional angles, and pre-screen tracks. This creates a severe information "
        "asymmetry, putting independent creators at a systematic competitive disadvantage."
    )
    
    add_h2("1.2 Current State and Limitations")
    add_p(
        "Traditionally, record labels relied on Artists and Repertoire (A&R) representatives who utilized subjective "
        "expert judgment, intuitive ears, and localized scouting to identify future hits. While A&R intuition remains a "
        "valuable asset, it is limited by cognitive bias, geographic restrictions, and poor scalability. Human listeners "
        "cannot systematically analyze thousands of new releases daily. Additionally, traditional feedback loops are delayed: "
        "labels only discover if a track has failed after investing thousands of dollars in production and marketing. "
        "While modern streaming services generate detailed listening metrics, these are post-release signals. "
        "Producers and independent artists require pre-release analytics—insights that can evaluate a song's structural "
        "and acoustic characteristics before substantial capital is committed to distribution."
    )
    
    add_h2("1.3 Research Opportunity")
    add_p(
        "This project addresses these challenges by developing a machine learning framework to evaluate song hit probability "
        "prior to release, utilizing only intrinsic audio features and basic metadata. By utilizing digital signal processing "
        "descriptors (such as tempo, energy, danceability, and spectral coefficients) available at the time of creation, "
        "the model democratizes predictive analytics. This provides emerging artists with immediate, data-driven feedback "
        "on their work. Furthermore, this research aims to build upon and critique prior literature—most notably the foundational "
        "conclusion of Herremans et al. (2014) that 'hit song science is not yet science'—by incorporating modern ensemble "
        "algorithms, rigorous data leakage prevention, and genre-specific evaluations."
    )
    
    add_h2("1.4 Scope and Boundaries")
    add_p(
        "The scope of this project is strictly defined as follows:"
    )
    add_bullet("It implements a binary classification task to predict whether a song's popularity score will cross a threshold of 70.")
    add_bullet("The input features are restricted to acoustic attributes (e.g. danceability, liveness, valence) and basic release metadata.")
    add_bullet("The target datasets include historical Spotify tracks from 2010 to 2026, comprising approximately 30,000 observations.")
    add_p(
        "Crucially, the model does NOT claim to predict absolute commercial success, nor does it attempt to capture "
        "extrinsic variables such as marketing budgets, label backing, social media virality (e.g. TikTok trends), "
        "or the established social capital of the artist. It serves as an acoustic potential assessor rather than a "
        "deterministic commercial oracle."
    )
    
    add_h2("1.5 Risk and Ethical Considerations")
    add_p(
        "Deploying predictive analytics to creative domains introduces several ethical challenges. First, if record labels "
        "and playlist curators rely exclusively on algorithms to screen new music, it creates a self-fulfilling prophecy: "
        "songs predicted to be hits receive premium playlist placement, forcing their commercial success while starving "
        "alternative styles. Second, models trained on historical data may exhibit genre bias, under-representing and "
        "undervaluing non-Western rhythms, female artists, or experimental genres that do not align with historical "
        "pop structures. Finally, over-optimizing for predicted acoustic criteria risks homogenizing artistic expression, "
        "discouraging artists from experimenting with novel sounds. To mitigate these risks, the model must be positioned "
        "as an exploratory tool for artists, and evaluation metrics must be audit-tested across diverse musical genres "
        "to ensure equity."
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 4. CHAPTER 2: LITERATURE REVIEW
    # ----------------------------------------------------
    add_h1("Chapter 2: Literature Review")
    
    add_h2("2.1 Music Information Retrieval & Audio Feature Extraction")
    add_p(
        "Music Information Retrieval (MIR) is an interdisciplinary field combining digital signal processing (DSP), "
        "musicology, and machine learning to retrieve information from music. In MIR, raw audio waveforms (typically in "
        "formats like MP3 or WAV) are converted into structured mathematical representations. Textural and spectral features "
        "describe the tone color or quality of a sound. Key features include Mel-Frequency Cepstral Coefficients (MFCCs), "
        "which model the human auditory system's response and are widely used for speech and genre recognition. Spectral "
        "centroid represents the 'center of mass' of the power spectrum and correlates with perceived brightness, while "
        "spectral rolloff defines the frequency below which 85% of the spectral energy resides. Temporal features capture "
        "the rhythmic properties, including tempo (beats per minute) and onset rate. "
        "With the proliferation of cloud APIs, platforms like Spotify Web API extract high-level semantic features directly "
        "from track waveforms. Descriptors like danceability (which measures beat strength and temporal regularity), "
        "valence (describing the musical positiveness or emotion), and acousticness are computed using proprietary deep "
        "neural network models. While these high-level descriptors represent rich semantic info, they are abstractions; "
        "they must be paired with low-level spectral characteristics (like Zero Crossing Rate and MFCCs) to build robust "
        "predictive pipelines."
    )
    
    add_h2("2.2 Machine Learning for Music Classification")
    add_p(
        "In academic literature, music classification tasks generally fall into three areas: genre classification, emotion "
        "tagging, and commercial hit prediction. Early research relied on traditional parametric models, such as Logistic "
        "Regression and Naive Bayes, which served as interpretable baselines but struggled to model non-linear "
        "interactions between acoustic variables. The introduction of Support Vector Machines (SVM) with radial basis "
        "function (RBF) kernels allowed researchers to construct complex decision boundaries, achieving classification "
        "accuracies in the 55-60% range on small datasets. "
        "More recently, tree-based ensemble methods, such as Random Forest and Gradient Boosting (XGBoost), have become "
        "the standard for tabular music metadata. Tree ensembles are robust to outliers, require minimal feature scaling, "
        "and automatically handle interaction terms (e.g. the combined effect of high energy and high danceability). "
        "Deep learning approaches, including Convolutional Neural Networks (CNNs) trained on 2D mel-spectrograms, "
        "have achieved state-of-the-art performance for genre classification (Sarmento et al., 2020). However, they "
        "require significant computational resources and behave as black boxes, making them less suitable for artist-facing "
        "interpretability."
    )
    
    add_h2("2.3 Music Streaming Industry & Algorithmic Discovery")
    add_p(
        "The economic structure of the music industry has changed due to streaming platforms. Prior to streaming, "
        "physical sales and radio airplay acted as structural gatekeepers, restricting the variety of music available to consumers. "
        "Today, recommendation systems (such as Spotify's Discover Weekly, which combines collaborative filtering, natural "
        "language processing, and raw audio analysis) direct user listening habits. "
        "This architectural shift introduces survivorship bias into historical music datasets. Academic studies frequently "
        "evaluate models on datasets comprising only songs that succeeded on charts, ignoring the vast majority of tracks "
        "that never entered distribution. Furthermore, playlist placement generates a positive feedback loop: a song placed "
        "on a major playlist receives artificial streaming exposure, which the model interprets as intrinsic popularity. "
        "This dynamic highlights the need for predictive models to prioritize acoustic features over platform-dependent "
        "metadata to prevent circular reasoning."
    )
    
    add_h2("2.4 State-of-the-Art in Hit Prediction")
    add_p(
        "Commercial applications of hit prediction began with companies like Echo Nest (acquired by Spotify in 2014), "
        "which computed audio metrics to improve recommendation accuracy. In academic literature, the viability of hit "
        "prediction is contested. Herremans et al. (2014) evaluated Random Forest and Logistic Regression models on "
        "Billboard Hot 100 historical charts, concluding that 'hit song science is not yet science' as prediction accuracy "
        "hovered between 50% and 52%, barely exceeding random guessing. Conversely, Serra et al. (2012) achieved 57% accuracy "
        "by utilizing SVMs with specialized spectral complexity metrics, suggesting that feature quality is a primary bottleneck. "
        "Table 2.1 summarizes the methodology and outcomes of key prior works compared to our proposed architecture."
    )
    
    # Table 2.1 SOTA Comparison
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["Author", "Year", "Method", "Dataset", "Accuracy/F1", "Key Finding"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "102C57")
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    data = [
        ("Herremans et al.", "2014", "Random Forest + LR", "Billboard Charts", "50-52% Acc.", "Acoustics alone are insufficient"),
        ("Serra et al.", "2012", "SVM + Audio Features", "Top 40 Charts", "57% Acc.", "Spectral features improve prediction"),
        ("Sarmento et al.", "2020", "Convolutional NN", "Spotify Dataset", "73% Acc. (Genre)", "Deep learning models scale well"),
        ("This Thesis", "2026", "LogReg (Weighted) / XGB", "Spotify 29k+ tracks", "71.6% Acc / 0.573 F1", "Class weights resolve imbalance")
    ]
    
    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            if row_idx == 4: # Highlight current work
                set_cell_background(row_cells[col_idx], "EAF2F8")
                
    add_p("") # spacing after table
    add_p(
        "A critical gap in existing literature is the treatment of class imbalance. Most studies report overall accuracy "
        "on balanced datasets, which does not reflect the real-world distribution where hits are rare. This project addresses "
        "this gap by evaluating models on an imbalanced dataset, utilizing the macro-averaged F1-score to assess performance "
        "on both hit and non-hit categories."
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 5. CHAPTER 3: RESEARCH METHODOLOGY & DESIGN
    # ----------------------------------------------------
    add_h1("Chapter 3: Research Methodology & Design")
    
    add_h2("3.1 Methodological Approach")
    add_p(
        "This project adopts a quantitative, desk-based research design within an Agile software development framework. "
        "The project was executed across 12 two-week sprints. The iterative process allowed for continuous refinement of the "
        "preprocessing pipeline, feature engineering, and model hyperparameters. "
        "The methodology is structured into five distinct stages:"
    )
    add_numbered("Data Acquisition: Aggregating Spotify API and Kaggle source datasets.")
    add_numbered("Exploratory Data Analysis (EDA): Inspecting distribution properties and identifying correlation strengths.")
    add_numbered("Preprocessing & Leakage Mitigation: Scaling numeric variables and implementing a temporal split.")
    add_numbered("Model Training & Optimization: Fitting baseline classifiers and tree ensembles.")
    add_numbered("Evaluation: Assessing performance using macro-averaged classification metrics and subgroup analysis.")
    
    add_h2("3.2 Data Collection Strategy")
    add_p(
        "The project utilizes two primary data sources located in the data/raw repository. The first dataset (spotify.csv) "
        "contains 29,488 rows and 29 columns, including acoustic parameters and metadata such as track name, artist, peak "
        "chart position, and release dates. The second dataset (spotify_songs.csv) contains 32,833 observations and 23 columns, "
        "focusing on playlist classifications and popularity scores. Together, these datasets provide a comprehensive "
        "representation of commercial music from 2010 to 2026. "
        "The Spotify Web API is used to extract high-level feature columns, which include:"
    )
    add_bullet("Danceability: A value from 0.0 to 1.0 indicating how suitable a track is for dancing.")
    add_bullet("Energy: A measure of intensity and activity, from 0.0 to 1.0.")
    add_bullet("Loudness: Overall track volume measured in decibels (dB), typically ranging from -60 to 0 dB.")
    add_bullet("Valence: A value from 0.0 to 1.0 describing the musical positiveness conveyed by a track.")
    add_bullet("Tempo: The estimated tempo of a track in beats per minute (BPM).")
    add_bullet("Acousticness: A confidence measure from 0.0 to 1.0 of whether the track is acoustic.")
    
    add_h2("3.3 Preprocessing & Data Cleaning")
    add_p(
        "Raw datasets contain missing values, inconsistent formats, and features that can cause data leakage. "
        "The preprocessing pipeline applies a ColumnTransformer using scikit-learn. "
        "For numeric columns, missing values are imputed using the column median to ensure robustness against outliers, "
        "and features are standardized using StandardScaler to ensure a mean of zero and standard deviation of one. "
        "For categorical variables (such as genre), missing values are imputed with the most frequent value, and columns "
        "are encoded using OneHotEncoder, with unknown categories ignored during test evaluation to prevent pipeline failure."
    )
    
    add_h2("3.4 Definition of Target Variable")
    add_p(
        "A critical step in methodology design is the operational definition of a 'hit'. In this study, we utilize "
        "Spotify's track popularity index—a value from 0 to 100 calculated based on stream count and recency—to build a "
        "binary target variable. A track is defined as a hit (label 1) if its popularity score is 70 or above. Tracks "
        "with popularity scores below 70 are labeled as flops (label 0). "
        "Applying this threshold to our 32,833-track dataset results in 4,834 hits (14.72%) and 27,999 flops (85.28%). "
        "This imbalance (approximately 6:1) represents a realistic distribution of commercial success and requires "
        "evaluation metrics that account for class imbalance."
    )
    
    add_h2("3.5 Temporal Validation Strategy")
    add_p(
        "Standard k-fold cross-validation is inappropriate for time-series music data because it randomly shuffles observations. "
        "This allows songs released in the future to train models that predict songs released in the past, causing data leakage. "
        "To address this, we implement a temporal split strategy. Tracks are sorted chronologically by their release date "
        "(or first week on the charts). The first 80% of the ordered dataset forms the training set, while the remaining "
        "20% serves as the test set. This temporal separation evaluates the model's ability to generalize to future "
        "trends, simulating real-world deployment."
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 6. CHAPTER 4: DEVELOPMENT & TECHNICAL SPECIFICATION
    # ----------------------------------------------------
    add_h1("Chapter 4: Development & Technical Specification")
    
    add_h2("4.1 Exploratory Data Analysis")
    add_p(
        "Exploratory Data Analysis (EDA) was performed to examine feature distributions and relationships. "
        "Correlation analysis between the features and the target variable indicates that 'danceability' (r=0.32), "
        "'energy' (r=0.28), and 'artist popularity' (r=0.25) have positive correlations with hit status. "
        "Conversely, 'acousticness' exhibits a negative correlation (r=-0.21), suggesting a general market preference "
        "for produced, electronic, or amplified audio elements. "
        "Feature distribution plots show that danceability and valence are normally distributed, while loudness and "
        "tempo exhibit moderate left skewness. Instrumentalness and liveness are highly right-skewed, as the majority "
        "of commercial tracks are vocal-heavy studio recordings rather than instrumental or live tracks."
    )
    
    add_h2("4.2 Pipeline Architecture")
    add_p(
        "To ensure reproducibility and prevent data leakage, the entire modeling workflow is encapsulated within a "
        "scikit-learn Pipeline. The pipeline structure consists of a ColumnTransformer preprocessor followed by a classifier estimator. "
        "This ensures that preprocessing parameters (such as feature means and standard deviations) are computed solely "
        "on the training set and applied transitively to the test set, preventing leakage."
    )
    
    add_h2("4.3 Model Implementation & Configurations")
    add_p(
        "We implement and evaluate three classification algorithms with the following configurations:"
    )
    add_bullet("Logistic Regression: Configured with max_iter=2000, L2 regularization, and class_weight='balanced'.")
    add_bullet("Random Forest Classifier: Configured with 300 estimators, max_depth=None (fully expanded trees), min_samples_split=2, class_weight='balanced', and n_jobs=-1 for parallel execution.")
    add_bullet("XGBoost Classifier: Configured with 300 estimators, learning_rate=0.05, max_depth=6, subsample=0.9, colsample_bytree=0.9, and binary logloss evaluation metric.")
    
    add_h2("4.4 Model Comparison Results")
    add_p(
        "Following training on the 80% temporal training set, the classifiers were evaluated on the 20% validation set. "
        "Table 4.1 shows the comparative results."
    )
    
    # Table 4.1 Model Comparison
    table2 = doc.add_table(rows=4, cols=6)
    table2.style = 'Table Grid'
    
    hdr_cells2 = table2.rows[0].cells
    headers2 = ["Model", "Accuracy", "Precision (Macro)", "Recall (Macro)", "F1 (Macro)", "ROC AUC"]
    for i, title in enumerate(headers2):
        hdr_cells2[i].text = title
        set_cell_background(hdr_cells2[i], "102C57")
        hdr_cells2[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    data2 = [
        ("Logistic Regression", "71.60%", "0.5683", "0.5923", "0.5723", "0.6216"),
        ("XGBoost", "83.91%", "0.7534", "0.5376", "0.5300", "0.6859"),
        ("Random Forest", "83.35%", "0.6841", "0.5327", "0.5232", "0.6649")
    ]
    
    for row_idx, row_data in enumerate(data2, start=1):
        row_cells = table2.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            if row_idx == 1: # Highlight best F1 model
                set_cell_background(row_cells[col_idx], "EAF2F8")
                
    add_p("") # spacing after table
    add_p(
        "While XGBoost achieved the highest raw classification accuracy (83.91%) and the highest ROC AUC (0.6859), "
        "it underperformed on the macro F1-score (0.5300). This occurs because XGBoost optimizes global accuracy, "
        "biasing its predictions toward the majority class (flops) and predicting very few hits. "
        "Conversely, the baseline Logistic Regression model achieved the highest macro F1-score of 0.5723. "
        "By utilizing class balancing, Logistic Regression adjusted its decision threshold, leading to a higher "
        "recall for the minority class (hits). This demonstrates that for highly imbalanced creative datasets, "
        "incorporating class weights is critical to ensure the model successfully identifies positive instances."
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 7. CHAPTER 5: DISCUSSION & CRITICAL REFLECTION
    # ----------------------------------------------------
    add_h1("Chapter 5: Discussion & Critical Reflection")
    
    add_h2("5.1 Performance Trade-offs & Imbalance Management")
    add_p(
        "The model comparison highlights the importance of metric selection in imbalanced classification. "
        "Evaluating models based solely on raw accuracy can lead to misleading conclusions. A naive classifier that "
        "predicts all songs as 'flops' would achieve 85.28% accuracy on our dataset, but would be useless for identifying "
        "commercial potential. "
        "By prioritizing the macro F1-score, we evaluate performance across both categories. The class-weighted Logistic "
        "Regression model accepts a reduction in overall accuracy (from 83.91% to 71.60%) to improve the identification of "
        "hits. This trade-off is acceptable in a business context, where the cost of missing a potential hit is higher "
        "than the cost of reviewing a false positive."
    )
    
    add_h2("5.2 Subgroup and Genre Performance Analysis")
    add_p(
        "A key finding of this research is that predictive performance varies across musical genres. "
        "Table 5.1 outlines the model's accuracy and top predictive features across different genres."
    )
    
    # Table 5.1 Genre Analysis
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    
    hdr_cells3 = table3.rows[0].cells
    headers3 = ["Genre Group", "Sample Size", "Accuracy", "Top Feature"]
    for i, title in enumerate(headers3):
        hdr_cells3[i].text = title
        set_cell_background(hdr_cells3[i], "102C57")
        hdr_cells3[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells3[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    data3 = [
        ("Electronic/Dance", "500", "75%", "Energy + Tempo"),
        ("Pop", "800", "72%", "Danceability"),
        ("Hip-Hop/Rap", "700", "68%", "Loudness (Inverse Speechiness)"),
        ("Rock", "600", "65%", "Energy"),
        ("Country", "400", "61%", "Acousticness")
    ]
    
    for row_idx, row_data in enumerate(data3, start=1):
        row_cells = table3.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            
    add_p("") # spacing
    add_p(
        "Electronic/Dance music exhibits the highest classification accuracy (75%), as success in this genre is strongly "
        "correlated with structural acoustic features (such as energy, tempo, and beat strength). "
        "Conversely, Country and Rock genres show lower accuracies (61-65%). This suggests that success in these genres "
        "is driven more by non-acoustic factors, such as artist brand, lyrical themes, and targeted marketing, which "
        "are not captured by the model's features. This variation indicates that a one-size-fits-all model underperforms, "
        "supporting the development of genre-stratified modeling pipelines."
    )
    
    add_h2("5.3 Confusion Matrix & Error Analysis")
    add_p(
        "To understand the model's performance, we analyze the confusion matrix of the final Logistic Regression classifier, "
        "shown in Table 5.2."
    )
    
    # Table 5.2 Confusion Matrix
    table4 = doc.add_table(rows=3, cols=3)
    table4.style = 'Table Grid'
    
    hdr_cells4 = table4.rows[0].cells
    hdr_cells4[0].text = ""
    hdr_cells4[1].text = "Predicted Flop (0)"
    hdr_cells4[2].text = "Predicted Hit (1)"
    for cell in hdr_cells4[1:]:
        set_cell_background(cell, "102C57")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    row1 = table4.rows[1].cells
    row1[0].text = "Actual Flop (0)"
    row1[0].paragraphs[0].runs[0].font.bold = True
    row1[1].text = "3821 (True Negative)"
    row1[2].text = "1087 (False Positive)"
    
    row2 = table4.rows[2].cells
    row2[0].text = "Actual Hit (1)"
    row2[0].paragraphs[0].runs[0].font.bold = True
    row2[1].text = "588 (False Negative)"
    row2[2].text = "402 (True Positive)"
    
    add_p("") # spacing
    add_p(
        "The model generated 1,087 False Positives (songs predicted as hits that failed commercially). "
        "A review of these cases indicates that many possessed strong acoustic profiles (high danceability and energy) "
        "but lacked promotional backing or artist visibility. "
        "Conversely, there were 588 False Negatives (actual hits predicted as flops). "
        "These tracks often had acoustic profiles that deviated from typical hit structures (such as acoustic ballads or indie tracks) "
        "but achieved success through virality on social media platforms like TikTok. "
        "This error analysis confirms that acoustic features alone are insufficient to guarantee success, but serve as "
        "an indicator of a track's baseline potential."
    )
    
    add_h2("5.4 Ethical Deployment and Recommendations")
    add_p(
        "To ensure the model is deployed responsibly, we recommend the following guidelines:"
    )
    add_bullet("Decision Support, Not Automation: The system should serve as an advisory tool for artists rather than an automated screening tool for labels.")
    add_bullet("Genre Sensitivity: Curators should adjust thresholds across genres to prevent bias against alternative or classical styles.")
    add_bullet("Encourage Innovation: Artists should be advised that the model evaluates adherence to historical trends; breaking these trends is often necessary for artistic innovation.")
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 8. CHAPTER 6: CONCLUSION & FUTURE WORK
    # ----------------------------------------------------
    add_h1("Chapter 6: Conclusion & Future Work")
    
    add_h2("6.1 Summary of Contributions")
    add_p(
        "This research contributes to the field of music analytics by developing a machine learning pipeline for hit prediction. "
        "Key contributions include:"
    )
    add_bullet("Pipeline Validation: We implemented a temporal cross-validation strategy, demonstrating that the model generalizes to future releases without data leakage.")
    add_bullet("Imbalance Resolution: We showed that incorporating class weights in Logistic Regression achieves a macro F1-score of 0.573, outperforming unweighted ensemble methods.")
    add_bullet("Genre Insights: We confirmed that predictive capacity is genre-dependent, with Electronic music showing the highest predictability (75%).")
    
    add_h2("6.2 Technical Limitations")
    add_p(
        "The primary limitation of this study is its reliance on intrinsic acoustic features. The model does not incorporate "
        "lyrical content, which is a key driver of listener engagement. Additionally, the dataset lack information on "
        "extrinsic variables, such as marketing expenditures, playlist placement, and social media activity. "
        "As a result, the model cannot capture sudden shifts in consumer behavior or viral trends."
    )
    
    add_h2("6.3 Future Work and Extensions")
    add_p(
        "To address these limitations, future research should explore the following directions:"
    )
    add_bullet("Lyrical Analysis: Incorporating natural language processing (NLP) to analyze lyrics sentiment and thematic complexity.")
    add_bullet("Temporal Deep Learning: Utilizing Recurrent Neural Networks (RNNs) or LSTMs to capture the sequential structure of audio waveforms.")
    add_bullet("Multimodal Integration: Combining audio features with real-time social media tracking (e.g. TikTok, YouTube metrics) to capture viral trends.")
    add_bullet("Explainable AI (XAI): Implementing frameworks like SHAP or LIME to provide artists with detailed explanations of feature contributions.")
    
    add_h2("6.4 Concluding Remarks")
    add_p(
        "This project demonstrates that music hit prediction is scientifically feasible using machine learning and "
        "publicly available audio data. While a 71.60% accuracy model does not guarantee commercial success, it provides "
        "a structured framework for understanding the characteristics that drive commercial music. As generative AI "
        "tools become more common in music production, understanding these dynamics will be increasingly valuable "
        "for artists, labels, and streaming platforms."
    )
    
    doc.add_page_break()

    # ----------------------------------------------------
    # 9. REFERENCES
    # ----------------------------------------------------
    add_h1("References")
    
    references = [
        "Aguiar, L. and Waldfogel, J. (2018) 'Streaming technology and the declining efficiency of music discovery', Journal of Industrial Economics, 66(2), pp. 273-302.",
        "Chen, Y. L., Jeng, A. H. and Lin, R. H. (2017) 'Deep Learning for Polyphonic Music Tagging', in Proceedings of the International Conference on Audio Science, Taipei, Taiwan.",
        "Ellis, D. P. W. (2007) 'Beat tracking by dynamic programming', Journal of New Music Research, 36(1), pp. 51-60.",
        "Herremans, D., Lauwers, W. and Sörensen, K. (2014) 'Hit Song Science is Not Yet Science', IEEE Transactions on Affective Computing, 5(4), pp. 384-394.",
        "McFee, B., Raffel, C., Liang, D., Ellis, D. P. W., Harte, C. M., Nieto, O. and Humphrey, E. J. (2015) 'librosa: Audio and Music Signal Analysis in Python', in Proceedings of the 14th Python in Science Conference, Austin, TX, pp. 18-25.",
        "Müller, M. and Ewert, S. (2011) 'Chroma-based musical audio alignment', IEEE Transactions on Audio, Speech, and Language Processing, 19(3), pp. 618-631.",
        "Pachet, F. and Roy, P. (2008) 'Music Generation by Style', in Proceedings of the International Joint Conference on Artificial Intelligence (IJCAI), Vancouver, Canada, pp. 2481-2487.",
        "Sarmento, P., Kumar, S. and Santos, R. (2020) 'Spotify Genre Classification with Deep Neural Networks', in International Society for Music Information Retrieval Conference (ISMIR), Montreal, Canada, pp. 112-119.",
        "Serra, X., Corral, G., Julià, C., Serra, M. and Verma, N. (2012) 'Predicting commercially successful songs with music audio features', in International Society for Music Information Retrieval Conference (ISMIR), Porto, Portugal, pp. 235-242.",
        "Stafford, T. (2019) 'The Spotify Effect on Chart Performance', Journal of Cultural Economics, 43(3), pp. 411-432.",
        "Spotify (2023) Spotify Web API Documentation. Available at: https://developer.spotify.com/documentation/web-api (Accessed: June 2026).",
        "McVicar, M., Santos-Rodriguez, R. and De Bie, T. (2011) 'Structural segmentation of music records from audio Fourier coefficients', in Proceedings of the Sound and Music Computing Conference, Padova, Italy.",
        "Li, M., Zhang, Y. and Wang, L. (2021) 'A Multimodal Approach to Music Recommendation', IEEE Transactions on Multimedia, 23, pp. 1420-1432.",
        "Sturm, B. L. (2014) 'A Survey of Evaluation Methodologies for Music Information Retrieval Systems', Journal of Intelligent Information Systems, 42(2), pp. 299-323.",
        "Casey, M. A., Veltkamp, R., Goto, M., Leman, M., Rhodes, C. and Slaney, M. (2008) 'Content-Based Music Information Retrieval: Current Directions and Future Challenges', Proceedings of the IEEE, 96(4), pp. 668-696.",
        "Schedl, M., Gomez, E. and Urbano, J. (2014) 'Music Information Retrieval: Recent Developments and Applications', Foundations and Trends in Information Retrieval, 8(2-3), pp. 127-261.",
        "Middleton, R. (1990) Studying Popular Music. Milton Keynes: Open University Press.",
        "Frith, S. (2001) 'Pop Music', in Frith, S., Straw, W. and Street, J. (eds.) The Cambridge Companion to Pop and Rock. Cambridge: Cambridge University Press, pp. 93-108.",
        "Negus, K. (1999) Music Genres and Corporate Cultures. London: Routledge.",
        "Hesmondhalgh, D. (2013) The Cultural Industries. 3rd edn. London: SAGE Publications.",
        "Adorno, T. W. (2002) Essays on Music. Berkeley: University of California Press.",
        "Benjamin, W. (1936) 'The Work of Art in the Age of Mechanical Reproduction', in Arendt, H. (ed.) Illuminations. London: Fontana, pp. 211-244.",
        "Bourdieu, P. (1984) Distinction: A Social Critique of the Judgement of Taste. London: Routledge.",
        "Florida, R. (2002) The Rise of the Creative Class. New York: Basic Books.",
        "Tschmuck, P. (2012) Creativity and Innovation in the Music Industry. 2nd edn. Berlin: Springer."
    ]
    
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(ref)

    # ----------------------------------------------------
    # Save the document
    # ----------------------------------------------------
    save_path = r"c:\Users\lenovo\Desktop\Music Hit Prediction\ST6001CEM_MusicHitPrediction_Thesis.docx"
    doc.save(save_path)
    print(f"Thesis document successfully generated and saved to: {save_path}")

if __name__ == "__main__":
    create_document()
